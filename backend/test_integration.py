import os
import requests
import docx

def create_sample_docx(filename="test_sample.docx"):
    print(f"[INFO] Creating sample DOCX: {filename}...")
    doc = docx.Document()
    
    # Add title and headings
    doc.add_heading("Quantum Mechanics and Ancient Philosophy", level=0)
    
    # Section 1: Theory in English
    doc.add_heading("Section 1: Quantum Superposition", level=1)
    doc.add_paragraph(
        "Quantum superposition is a fundamental principle of quantum mechanics. It states that, "
        "much like waves in classical physics, any two (or more) quantum states can be added together "
        "('superposed') and the result will be another valid quantum state. Conversely, that every quantum "
        "state can be represented as a sum of two or more other distinct states. Mathematically, it refers "
        "to a property of solutions to the Schrödinger equation; since the Schrödinger equation is linear, "
        "any linear combination of solutions will also be a solution."
    )
    doc.add_paragraph(
        "Mathematically, if |psi1> and |psi2> are valid states, then a superposed state is given by: "
        "|psi> = a|psi1> + b|psi2>, where a and b are complex probability amplitudes such that |a|^2 + |b|^2 = 1."
    )
    
    # Section 2: Math Exam Paper style
    doc.add_heading("Section 2: Practice Mathematics Problem", level=1)
    doc.add_paragraph(
        "Solve the following quadratic equation using the quadratic formula: "
        "x^2 - 5x + 6 = 0."
    )
    doc.add_paragraph(
        "Hint: The quadratic formula is x = (-b +- sqrt(b^2 - 4ac)) / (2a). "
        "For this equation, a = 1, b = -5, and c = 6. The roots are x = 2 and x = 3."
    )

    # Section 3: Sanskrit Shloka and Hindi Translation
    doc.add_heading("Section 3: Sanskrit Wisdom", level=1)
    doc.add_paragraph(
        "कर्मण्येवाधिकारस्ते मा फलेषु कदाचन।\n"
        "मा कर्मफलहेतुर्भूर्मा ते सङ्गोऽस्त्वकर्मणि॥"
    )
    doc.add_paragraph(
        "Transliteration: Karmanve vadhikaraste ma phaleshu kadachana, ma karmaphalaheturbhurma te sango'stvakarmani."
    )
    doc.add_paragraph(
        "Hindi Meaning: तुम्हारा अधिकार केवल कर्म करने पर है, उसके फलों पर कभी नहीं। "
        "इसलिए तुम कर्मों के फल की इच्छा वाले मत बनो और तुम्हारी कर्म न करने में भी आसक्ति न हो।"
    )
    
    doc.save(filename)
    print("[SUCCESS] Sample DOCX created.")

def run_integration_test():
    base_url = "http://localhost:8000"
    filename = "test_sample.docx"
    
    # 1. Create document
    create_sample_docx(filename)
    
    # 2. Upload document
    print("\n[INFO] Uploading document to /upload...")
    with open(filename, "rb") as f:
        files = {"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = requests.post(f"{base_url}/upload", files=files)
        
    if resp.status_code != 200:
        print(f"[ERROR] Upload failed with status {resp.status_code}: {resp.text}")
        return
        
    data = resp.json()
    session_id = data["session_id"]
    print(f"[SUCCESS] Upload succeeded! Session ID: {session_id}")
    print(f"[INFO] Parsed details: Type={data['file_type']}, Pages={data['page_count']}, Words={data['word_count']}")
    
    # 3. Chat test 1: English question about Quantum Superposition
    print("\n[INFO] Chat Test 1: Asking in English about Superposition...")
    chat_resp = requests.post(f"{base_url}/chat", json={
        "session_id": session_id,
        "question": "What is quantum superposition and what is the mathematical formula for it?"
    })
    
    if chat_resp.status_code != 200:
        print(f"[ERROR] Chat Test 1 failed with status {chat_resp.status_code}: {chat_resp.text}")
    else:
        chat_data = chat_resp.json()
        print(f"[AI] Answer (English):\n{chat_data['answer']}\n")
        print(f"[INFO] Sources: {chat_data['sources']}")
        
    # 4. Chat test 2: Hindi question about Sanskrit Shloka
    print("\n[INFO] Chat Test 2: Asking in Hindi about the Sanskrit shloka...")
    chat_resp_2 = requests.post(f"{base_url}/chat", json={
        "session_id": session_id,
        "question": "इस दस्तावेज़ में दिए गए संस्कृत श्लोक का अर्थ और उसका दार्शनिक महत्व क्या है?"
    })
    
    if chat_resp_2.status_code != 200:
        print(f"[ERROR] Chat Test 2 failed with status {chat_resp_2.status_code}: {chat_resp_2.text}")
    else:
        chat_data_2 = chat_resp_2.json()
        # Since Hindi might also trigger console encoding issues in print if output is directed to a standard CP1252 console,
        # we try printing, and fallback if it fails.
        try:
            print(f"[AI] Answer (Hindi):\n{chat_data_2['answer']}\n")
        except UnicodeEncodeError:
            print("[INFO] Answer (Hindi) received successfully, but cannot be printed due to console encoding limits.")
            # Print an ascii representation or safe string
            print(f"[AI] Answer (Hindi, safe bytes): {chat_data_2['answer'].encode('ascii', 'ignore').decode('ascii')}")
        
    # 5. Clean up
    print("\n[INFO] Cleaning up session...")
    del_resp = requests.delete(f"{base_url}/session/{session_id}")
    if del_resp.status_code == 200:
        print("[SUCCESS] Session deleted successfully.")
    else:
        print(f"[ERROR] Failed to delete session: {del_resp.text}")
        
    # Remove local test file
    if os.path.exists(filename):
        os.remove(filename)
        print("[SUCCESS] Local test file removed.")

if __name__ == "__main__":
    try:
        run_integration_test()
    except Exception as e:
        print(f"[ERROR] Integration test encountered an error: {str(e)}")
