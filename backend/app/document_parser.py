"""
Document parser — extracts text from PDF and DOCX files.
Supports multilingual content including Hindi, English, and Sanskrit.
"""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> dict:
    """
    Extract text and metadata from a PDF file.
    Uses pypdf, pdfplumber fallback, and page-by-page LLaMA 3.2 Vision OCR for sparse pages.
    """
    from pypdf import PdfReader
    import pdfplumber
    import pypdfium2 as pdfium
    import base64
    import time
    from io import BytesIO
    from groq import Groq
    from app.config import settings

    result = {
        "text": "",
        "pages": [],
        "metadata": {},
        "page_count": 0,
        "file_type": "pdf",
    }

    try:
        # Primary extraction with pypdf
        reader = PdfReader(io.BytesIO(file_bytes))
        result["page_count"] = len(reader.pages)

        # Extract metadata
        meta = reader.metadata
        if meta:
            result["metadata"] = {
                "title": meta.title or "",
                "author": meta.author or "",
                "subject": meta.subject or "",
                "creator": meta.creator or "",
            }

        # Initialize Groq client if key is configured
        client = None
        if settings.GROQ_API_KEY:
            client = Groq(api_key=settings.GROQ_API_KEY)
        else:
            logger.warning("GROQ_API_KEY is missing. Skipping LLaMA Vision OCR.")

        ocr_pages_count = 0
        max_ocr_pages = 50   # Process up to 50 sparse/image pages per document
        ocr_fatal_error = False  # If a fatal API error occurs, abort remaining OCR attempts

        # We'll use pdfplumber on-demand for pages that are sparse in pypdf
        try:
            plumber_pdf = pdfplumber.open(io.BytesIO(file_bytes))
        except Exception as e:
            logger.warning(f"Could not open pdfplumber: {e}")
            plumber_pdf = None

        try:
            pdfium_doc = pdfium.PdfDocument(file_bytes)
        except Exception as e:
            logger.warning(f"Could not open pypdfium2: {e}")
            pdfium_doc = None

        ocr_processed_any = False

        for page_num in range(result["page_count"]):
            # 1. Primary pypdf extraction
            page_text = ""
            try:
                page_text = reader.pages[page_num].extract_text() or ""
            except Exception as e:
                logger.warning(f"pypdf extraction failed on page {page_num + 1}: {e}")

            # 2. Try pdfplumber fallback for tables/text if pypdf is thin (< 150 chars)
            if len(page_text.strip()) < 150 and plumber_pdf and page_num < len(plumber_pdf.pages):
                try:
                    plumber_page = plumber_pdf.pages[page_num]
                    extracted = plumber_page.extract_text()
                    if extracted and len(extracted.strip()) > len(page_text.strip()):
                        page_text = extracted

                    # Extract tables
                    tables = plumber_page.extract_tables()
                    if tables:
                        table_text = ""
                        for table in tables:
                            for row in table:
                                row_text = " | ".join([cell or "" for cell in row])
                                table_text += row_text + "\n"
                        if len(table_text.strip()) > 0:
                            page_text += "\n" + table_text
                except Exception as e:
                    logger.warning(f"pdfplumber failed on page {page_num + 1}: {e}")

            # 3. Only OCR pages that are truly blank/image-only (< 50 chars after all extractors).
            #    Pages with any real text content (>= 50 chars) skip OCR entirely — faster & more accurate.
            if len(page_text.strip()) < 50:
                if ocr_fatal_error:
                    page_text = f"[Page {page_num + 1}: OCR disabled due to a previous fatal API error.]"
                elif client and pdfium_doc and page_num < len(pdfium_doc):
                    if ocr_pages_count < max_ocr_pages:
                        logger.info(f"🚨 Sparse/Scanned page detected (Page {page_num + 1}). Running Vision OCR...")
                        try:
                            # Render page to PIL image (scale=1.5 keeps base64 size under 4MB limit)
                            page = pdfium_doc[page_num]
                            bitmap = page.render(scale=1.5)
                            pil_img = bitmap.to_pil()

                            # Convert to base64 JPEG (quality=70 keeps size well under 4MB)
                            buffered = BytesIO()
                            pil_img.save(buffered, format="JPEG", quality=70)
                            img_bytes = buffered.getvalue()
                            img_str = base64.b64encode(img_bytes).decode("utf-8")
                            logger.info(f"  Image size: {len(img_bytes) / 1024:.1f} KB")

                            # Run Multimodal OCR with Llama 4 Scout
                            response = client.chat.completions.create(
                                model="meta-llama/llama-4-scout-17b-16e-instruct",
                                messages=[
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": (
                                                    "Extract all text from this page. Preserve the layout, headings, tables, "
                                                    "and math equations exactly as they appear in the image. Use standard LaTeX "
                                                    "for all math and formulas. Respond only with the extracted text, do not add "
                                                    "any introduction, explanation, or side notes."
                                                )
                                            },
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:image/jpeg;base64,{img_str}"
                                                }
                                            }
                                        ]
                                    }
                                ],
                                temperature=0.1,
                                max_tokens=2048
                            )
                            ocr_result = response.choices[0].message.content or ""
                            if len(ocr_result.strip()) > 10:
                                page_text = ocr_result
                                ocr_pages_count += 1
                                ocr_processed_any = True
                                logger.info(f"🎉 Vision OCR completed for Page {page_num + 1}.")
                            
                            # Pause 1s between OCR requests to avoid rate limits
                            time.sleep(1)
                        except Exception as page_err:
                            err_str = str(page_err)
                            logger.error(f"Failed OCR on page {page_num + 1}: {page_err}")
                            # Detect fatal errors (decommissioned model, auth, etc.) — abort immediately
                            if any(code in err_str for code in ["model_decommissioned", "invalid_request_error", "authentication", "401", "403"]):
                                logger.error(f"❌ Fatal OCR error detected — disabling OCR for remaining pages.")
                                ocr_fatal_error = True
                                page_text = f"[OCR unavailable for Page {page_num + 1}: fatal API error — {err_str[:120]}]"
                            else:
                                page_text = f"[OCR Error on Page {page_num + 1}: {err_str}]"
                    else:
                        logger.warning(f"Page {page_num + 1} is sparse but OCR limit of {max_ocr_pages} has been reached.")
                        page_text = f"[Page {page_num + 1} was not processed by Vision OCR — document exceeds the 50-page OCR limit.]"
                else:
                    if not client:
                        logger.warning(f"Page {page_num + 1} is sparse, but GROQ_API_KEY is missing. OCR skipped.")

            # Store result
            result["pages"].append({
                "page_number": page_num + 1,
                "text": page_text,
            })
            result["text"] += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

        # Clean up open resources
        if plumber_pdf:
            try:
                plumber_pdf.close()
            except Exception:
                pass

        if ocr_processed_any:
            result["metadata"]["ocr_processed"] = True

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    return result


def extract_text_from_docx(file_bytes: bytes) -> dict:
    """
    Extract text and metadata from a DOCX file.
    Preserves paragraph structure, headings, and tables.
    """
    from docx import Document

    result = {
        "text": "",
        "pages": [],
        "metadata": {},
        "page_count": 1,
        "file_type": "docx",
    }

    try:
        doc = Document(io.BytesIO(file_bytes))

        # Extract core properties
        props = doc.core_properties
        result["metadata"] = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
        }

        # Extract paragraphs with style information
        full_text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else "Normal"
                if "Heading" in style:
                    full_text_parts.append(f"\n## {para.text}\n")
                else:
                    full_text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = "\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(cells) + "\n"
            full_text_parts.append(table_text)

        result["text"] = "\n".join(full_text_parts)
        result["pages"] = [{"page_number": 1, "text": result["text"]}]

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    return result


def parse_document(file_bytes: bytes, filename: str) -> dict:
    """
    Auto-detect file type and extract text content.
    Returns a unified document structure.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        doc_data = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        doc_data = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type. Please upload a PDF or DOCX file.")

    doc_data["filename"] = filename
    doc_data["char_count"] = len(doc_data["text"])
    doc_data["word_count"] = len(doc_data["text"].split())

    if doc_data["word_count"] < 10:
        from app.config import settings
        if doc_data["file_type"] == "pdf" and not settings.GROQ_API_KEY:
            raise ValueError(
                "Could not extract meaningful text from this PDF document. "
                "It appears to contain scanned images/formulas with no embedded text, and "
                "GROQ_API_KEY is not configured in backend/.env to run LLaMA Vision OCR. "
                "Please configure your API key or upload a text-based PDF/DOCX file."
            )
        raise ValueError(
            "Could not extract meaningful text from the document. "
            "The file may be a scanned image or corrupted. "
            "Please upload a text-based PDF or DOCX file."
        )

    logger.info(
        f"Parsed document: {filename}, "
        f"{doc_data['page_count']} pages, "
        f"{doc_data['word_count']} words"
    )

    return doc_data
